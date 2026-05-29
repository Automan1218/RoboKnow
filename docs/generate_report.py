#!/usr/bin/env python3
"""
Generate PaiSmart_Project_Report.docx from project-report.md.
Corrections applied:
  - Solo developer (removes team/multi-member references)
  - No K8s (ADR-07 updated; only future-enhancement section keeps scaling notes)
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def strip_md(text):
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text.strip()


def add_runs(para, text, default_bold=False):
    """Add text to paragraph with inline **bold** and `code` formatting."""
    if not text:
        return
    parts = re.split(r'(\*\*[^*]+?\*\*|`[^`]+?`|\*[^*]+?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            run = para.add_run(part)
            if default_bold:
                run.bold = True


def is_separator_row(line):
    return bool(re.match(r'^\|[\-: |]+\|$', line.strip()))


def parse_md_table(table_lines):
    rows = []
    for line in table_lines:
        if is_separator_row(line):
            continue
        parts = line.split('|')
        cells = [c.strip() for c in parts[1:-1]]
        if cells:
            rows.append(cells)
    return rows


def add_code_block(doc, code_lines):
    """Add code lines as monospaced paragraphs with gray background."""
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run(line if line.strip() else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)
    # Trailing spacing after the block
    end = doc.add_paragraph()
    end.paragraph_format.space_before = Pt(0)
    end.paragraph_format.space_after = Pt(6)


def add_md_table(doc, rows):
    """Add parsed markdown table to the document."""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < num_cols:
            r.append('')

    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for ri, row in enumerate(rows):
        is_header = (ri == 0)
        for ci in range(num_cols):
            cell_text = row[ci] if ci < len(row) else ''
            cell = table.rows[ri].cells[ci]
            original = cell_text.strip()
            is_bold = is_header or bool(re.match(r'^\*\*[^*]*\*\*$', original))
            plain = strip_md(original)

            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(plain)
            run.font.size = Pt(10)
            if is_bold and plain:
                run.bold = True
            if is_header:
                set_cell_bg(cell, 'D9E1F2')

    spacing = doc.add_paragraph()
    spacing.paragraph_format.space_before = Pt(0)
    spacing.paragraph_format.space_after = Pt(6)


def add_page_numbers(doc):
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'end')
    run._r.append(fc1)
    run._r.append(instr)
    run._r.append(fc2)


# ─── Content Corrections ──────────────────────────────────────────────────────

def apply_corrections(content):
    # ADR-01: team → solo developer
    content = content.replace(
        'Team of ≤4 developers; single EC2 deployment target; '
        'no requirement for independent scaling of individual services at this stage.',
        'Solo developer project; single EC2 deployment target; '
        'no requirement for independent scaling of individual services at this stage.'
    )

    # Stakeholder register: Development Team → Developer (Solo)
    content = content.replace(
        '| Development Team | Internal | Build, maintain, deploy | High |',
        '| Developer (Solo) | Internal | Single developer responsible for all '
        'backend, frontend, AI integration, DevOps, testing, and deployment | High |'
    )

    # ADR-07: remove K8s migration reference
    content = content.replace(
        'K8s migration path documented (see Section 23).',
        'Horizontal scaling path is documented in Section 23 (Future Enhancements).'
    )

    # Section 8.2: replace multi-member table with solo developer breakdown
    old_82 = (
        "### 8.2 Per-Member Effort Breakdown\n"
        "\n"
        "| Member | Role | Sprint 1 | Sprint 2 | Sprint 3 | Sprint 4 | Total |\n"
        "|--------|------|----------|----------|----------|----------|-------|\n"
        "| [Name 1] | Backend Lead | – | – | – | – | – |\n"
        "| [Name 2] | Frontend Lead | – | – | – | – | – |\n"
        "| [Name 3] | DevOps / QA | – | – | – | – | – |\n"
        "| [Name 4] | Full Stack | – | – | – | – | – |\n"
        "| **Total** | | | | | | |\n"
        "\n"
        "> Fill in actual hours from your project management tool "
        "(GitHub Projects / Jira / Trello)."
    )
    new_82 = (
        "### 8.2 Solo Developer Effort Breakdown\n"
        "\n"
        "This project was built entirely by a single developer. "
        "All backend, frontend, AI integration, infrastructure, testing, and "
        "deployment work was performed individually.\n"
        "\n"
        "| Sprint | Theme | Estimated (hrs) | Actual (hrs) | Key Challenge |\n"
        "|--------|-------|----------------|--------------|---------------|\n"
        "| Sprint 1 | Foundation | 40 | ~45 | "
        "Docker Compose startup-order issue; custom ES IK plugin image |\n"
        "| Sprint 2 | RAG Core | 60 | ~65 | "
        "ES dense_vector dimension mismatch required full index recreation |\n"
        "| Sprint 3 | Agent Intelligence | 70 | ~80 | "
        "STM/LTM design iterated twice; new WebSocket event types for agent UI |\n"
        "| Sprint 4 | Quality & Delivery | 50 | ~55 | "
        "CI Elasticsearch healthcheck flakiness; OpenAI provider migration |\n"
        "| **Total** | **30 user stories** | **220** | **~245** | "
        "**Solo effort; +11% overall variance** |"
    )
    content = content.replace(old_82, new_82)

    return content


# ─── Markdown → Docx processor ────────────────────────────────────────────────

def process(doc, content):
    lines = content.split('\n')
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Empty line
        if not stripped:
            i += 1
            continue

        # Horizontal rule — skip (headings provide visual breaks)
        if stripped == '---':
            i += 1
            continue

        # Title (H1 — only the document title)
        if re.match(r'^# [^#]', line):
            h = doc.add_heading(line[2:].strip(), level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # H2
        if re.match(r'^## [^#]', line):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue

        # H3
        if re.match(r'^### [^#]', line):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        # H4
        if re.match(r'^#### [^#]', line):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        # Code block
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < n:
                i += 1  # skip closing ```
            add_code_block(doc, code_lines)
            continue

        # Table
        if stripped.startswith('|'):
            tbl_lines = []
            while i < n and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i].strip())
                i += 1
            rows = parse_md_table(tbl_lines)
            add_md_table(doc, rows)
            continue

        # Blockquote
        if line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(strip_md(text))
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Diagram placeholder  `[DIAGRAM — ...]`
        if stripped.startswith('`[DIAGRAM'):
            text = stripped.strip('`')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.2)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'FFF3CD')
            pPr.append(shd)
            run = p.add_run(text)
            run.italic = True
            run.bold = True
            run.font.color.rgb = RGBColor(0x85, 0x65, 0x00)
            i += 1
            continue

        # Unordered list (top-level: starts with "- " or "* ")
        if re.match(r'^[-*] ', line):
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, text)
            i += 1
            continue

        # Unordered list (indented 2–4 spaces)
        if re.match(r'^ {2,4}[-*] ', line):
            text = re.sub(r'^ +[-*] ', '', line).strip()
            p = doc.add_paragraph(style='List Bullet 2')
            add_runs(p, text)
            i += 1
            continue

        # Ordered list (top-level)
        if re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            add_runs(p, text)
            i += 1
            continue

        # Ordered list (indented)
        if re.match(r'^ {2,}\d+\. ', line):
            text = re.sub(r'^ +\d+\. ', '', line).strip()
            p = doc.add_paragraph(style='List Number 2')
            add_runs(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1


# ─── Entry point ──────────────────────────────────────────────────────────────

def main():
    base = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(base, 'project-report.md')
    out_path = os.path.join(base, 'PaiSmart_Project_Report.docx')

    print(f'Reading  : {md_path}')
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    print('Applying corrections...')
    content = apply_corrections(content)

    print('Building Word document...')
    doc = Document()

    # Page margins
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1.25)
        s.right_margin = Inches(1.25)

    # Base font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # Footer: page numbers
    add_page_numbers(doc)

    process(doc, content)

    doc.save(out_path)
    print(f'Saved    : {out_path}')
    print('Done.')


if __name__ == '__main__':
    main()
